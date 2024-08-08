from selenium.common.exceptions import NoSuchElementException, ElementNotVisibleException, ElementNotSelectableException
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from datetime import datetime
import pyautogui
import logging
from docx import Document
from docx.shared import Cm, RGBColor, Pt
import subprocess
import os
from dotenv import load_dotenv

# Load environment variables from the .env file
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def start_driver():
    chrome_options = Options()
    arguments = ['--lang=en-GB', '--window-size=1300,1000', '--incognito']
    for argument in arguments:
        chrome_options.add_argument(argument)

    chrome_options.add_experimental_option('prefs', {
        'download.prompt_for_download': False,
        'profile.default_content_setting_values.notifications': 2,
        'profile.default_content_setting_values.automatic_downloads': 1
    })
    driver = webdriver.Chrome(options=chrome_options)

    wait = WebDriverWait(
        driver,
        10,
        poll_frequency=1,
        ignored_exceptions=[NoSuchElementException, ElementNotVisibleException, ElementNotSelectableException]
    )
    return driver, wait

def get_dollar_value_and_date(driver, wait, desktop_path):
    logging.info('Accessing the site')
    driver.get('https://www.xe.com/pt/currencyconverter/convert/?Amount=1&From=USD&To=BRL')

    try:
        logging.info('Waiting for the page to load')
        wait.until(EC.presence_of_element_located((By.XPATH, "//p[@class='sc-e08d6cef-1 fwpLse'][1]")))

        screenshot_path = os.path.join(desktop_path, 'screen.jpg')
        logging.info('Taking a screenshot')
        pyautogui.screenshot(screenshot_path)
        logging.info('Screenshot taken successfully!')

        logging.info('Retrieving the current Dollar value')
        dollar_element = driver.find_element(By.XPATH, "//p[@class='sc-e08d6cef-1 fwpLse'][1]")
        text = dollar_element.text
        dollar = text.split()[0][:4]
        date = datetime.now().strftime('%d/%m/%Y  —  %H:%M')
        
    except NoSuchElementException:
        logging.error('Dollar value not found')
        return None, None
    except Exception as e:
        logging.error(f'Error retrieving dollar value: {e}')
        return None, None
    finally:
        driver.quit()
    return dollar, date, screenshot_path

def add_info_to_word_file(dollar, date, screenshot_path, desktop_path):
    logging.info('Creating a Word document')
    document = Document()
    logging.info('Formatting the Word document with collected information')

    title = document.add_heading(f'Current Dollar Rate — {dollar} ({date})', level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(20)

    document.add_paragraph('')
   
    paragraph = document.add_paragraph()
    paragraph.add_run(f'The dollar is valued at {dollar}, on {date}. ').font.color.rgb = RGBColor(0, 0, 0)

    document.add_paragraph('Quoted value on the site:  https://www.xe.com/pt/currencyconverter/convert/?Amount=1&From=USD&To=BRL')
    
    title2 = document.add_heading('Current Rate Screenshot', level=1)
    for run in title2.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    document.add_picture(screenshot_path, width=Cm(15.25))

    final_paragraph = document.add_paragraph()
    run = final_paragraph.add_run('Quotation made by – Patrícia Jorge')
    run.font.size = Pt(12)
    
    docx_path = os.path.join(desktop_path, 'Dollar_Rate_Report.docx')
    logging.info('Saving the Word document')
    document.save(docx_path)
    return docx_path

def convert_docx_to_pdf(docx_path, desktop_path):
    logging.info('Converting DOCX to PDF')
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"The file {docx_path} was not found.")
    
    try:
        soffice_path = os.getenv('SOFFICE_PATH')
        if not soffice_path:
            raise EnvironmentError("SOFFICE_PATH environment variable not set.")

        subprocess.call([soffice_path, '--headless', '--convert-to', 'pdf', '--outdir', desktop_path, docx_path])
        
        pdf_path = os.path.join(desktop_path, os.path.basename(docx_path).replace('.docx', '.pdf'))
        logging.info("File successfully converted!")
        return pdf_path
    except Exception as e:
        logging.error(f"Error converting the file: {e}")
        return None

if __name__ == '__main__':
    driver, wait = start_driver()
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    dollar, date, screenshot_path = get_dollar_value_and_date(driver, wait, desktop_path)
    
    if dollar and date:
        docx_file = add_info_to_word_file(dollar, date, screenshot_path, desktop_path)
        pdf_file = convert_docx_to_pdf(docx_file, desktop_path)
        
        if docx_file and pdf_file and screenshot_path:
            logging.info('All files created on desktop')
            
    logging.info('Automation complete!')
